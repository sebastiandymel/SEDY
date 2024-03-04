import docx
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor

def read_file(filename):
    with open(filename, 'r') as file:
        lines = file.readlines()
    return lines

def extract_sections(file_lines):
    sections = {}
    current_section = None
    for line in file_lines:
        if line.startswith("###"):
            current_section = line.strip()
            sections[current_section] = []
        elif current_section:
            sections[current_section].append(line.strip())
    return sections

def add_black_borders_to_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Set font size to 10 Pt
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black

def generate_comparison_report(file1, file2, output_file):
    doc = docx.Document()

    sections1 = extract_sections(read_file(file1))
    sections2 = extract_sections(read_file(file2))

    for section in sections1:
        if section in sections2:
            section_name = section.strip("###").strip()
            doc.add_heading(section_name, level=1)

            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False

            table.cell(0, 0).text = file1
            table.cell(0, 1).text = file2
            table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

            common_items = set(sections1[section]) & set(sections2[section])

            for item in common_items:
                row = table.add_row().cells
                cell1, cell2 = row[0], row[1]
                cell1.text = item
                cell2.text = item

            # Add items that exist in file1 but not in file2
            for item in sections1[section]:
                if item not in common_items:
                    row = table.add_row().cells
                    cell1, cell2 = row[0], row[1]
                    cell1.text = item
                    for paragraph in cell1.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)  # Set font size to 10 Pt
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Red

            # Add items that exist in file2 but not in file1
            for item in sections2[section]:
                if item not in common_items:
                    row = table.add_row().cells
                    cell1, cell2 = row[0], row[1]
                    cell2.text = item
                    for paragraph in cell2.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)  # Set font size to 10 Pt
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Red

            # Add black borders and adjust font size
            add_black_borders_to_table(table)

            doc.add_page_break()

    doc.save(output_file)

file1 = "2023-10-25-output.txt"
file2 = "2022-08-29-output.txt"
output_file = "comparison_report.docx"
generate_comparison_report(file1, file2, output_file)
